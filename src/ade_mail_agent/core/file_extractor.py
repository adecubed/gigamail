"""
file_extractor.py — Estrae testo da PDF, Word, Excel, TXT.
"""

import os
from typing import Tuple


def extract_text(file_path: str, original_filename: str = '') -> Tuple[str, str]:
    """
    Estrae testo da un file.
    Ritorna (testo_estratto, tipo_file).
    """
    ext = os.path.splitext(original_filename or file_path)[1].lower()

    if ext == '.pdf':
        return _extract_pdf(file_path), 'pdf'
    elif ext in ('.docx', '.doc'):
        return _extract_word(file_path), 'word'
    elif ext in ('.xlsx', '.xls'):
        return _extract_excel(file_path), 'excel'
    elif ext in ('.txt', '.md', '.csv'):
        return _extract_text(file_path), 'text'
    else:
        # Prova come testo
        try:
            return _extract_text(file_path), 'text'
        except Exception:
            raise ValueError(f'Formato non supportato: {ext}')


def _extract_pdf(path: str) -> str:
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return '\n'.join(text).strip()
    except ImportError:
        raise ValueError('pdfplumber non installato. Esegui: pip install pdfplumber')


def _extract_word(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Includi anche testo dalle tabelle
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return '\n'.join(paragraphs).strip()


def _extract_excel(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f'[Foglio: {sheet.title}]')
        for row in sheet.iter_rows(values_only=True):
            row_text = ' | '.join(str(v) for v in row if v is not None)
            if row_text.strip():
                lines.append(row_text)
    return '\n'.join(lines).strip()


def _extract_text(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read().strip()